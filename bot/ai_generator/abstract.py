import io
import re

from docx import Document
from docx.shared import Inches

try:
    from image_scrapper import downloader
except ImportError:
    from .image_scrapper import downloader


async def generate_docx_prompt(language, emotion_type, topic):
    message = f"""Create an {language} language very long outline for a {emotion_type} research paper on the topic of {topic} which is as comprehensive as possible. 
Language of research paper - {language}.
Provide in-depth and detailed information on each aspect.

Put this tag before the Title: [TITLE]
Put this tag after the Title: [/TITLE]
Put this tag before the Subtitle: [SUBTITLE]
Put this tag after the Subtitle: [/SUBTITLE]
Put this tag before the Heading: [HEADING]
Put this tag after the Heading: [/HEADING]
Put this tag before the Content: [CONTENT]
Put this tag after the Content: [/CONTENT]
Put this tag before the Image: [IMAGE]
Put this tag after the Image: [/IMAGE]

Elaborate extensively on the Content, ensuring a thorough exploration of each topic.
Conclude each Content section with [/CONTENT].

For instance:
[TITLE]Mental Health[/TITLE]
[SUBTITLE]Understanding and Nurturing Your Mind: A Comprehensive Guide to Mental Health[/SUBTITLE]
[HEADING]Mental Health Definition[/HEADING]
[CONTENT]...[/CONTENT]
[IMAGE]Person Meditating[/IMAGE]

Pay meticulous attention to the language of the research paper - {language}.
Accompany each image with descriptive keywords, such as "Mount Everest Sunset" or "Niagara Falls Rainbow".
Avoid discussing the research paper itself in the response (e.g., "Include pictures here about...").
Ensure the Title remains free of any special characters (?, !, ., :, ).
Strictly adhere to the specified format without including any additional information."""



    return message


async def generate_docx(answer):
    doc = Document()

    async def split_tags(reply):
        pattern = r'\[(.*?)\](.*?)\[/\1\]'
        tags = re.findall(pattern, reply, re.DOTALL)
        return tags

    async def parse_response(tags_array):
        if not tags_array:
            raise IndexError
        for item in tags_array:
            match (item[0]):
                case('TITLE'):
                    doc.add_heading(item[1], 0)
                case('SUBTITLE'):
                    doc.add_heading(item[1], 1)
                case('HEADING'):
                    doc.add_heading(item[1], 2)
                case('CONTENT'):
                    doc.add_paragraph(item[1])
                case('IMAGE'):
                    try:
                        image_data = await downloader.download(item[1], limit=1, adult_filter_off=True, timeout=15,
                                                               filter="+filterui:aspect-wide+filterui:imagesize-wallpaper+filterui:photo-photo")
                        doc.add_picture(io.BytesIO(image_data), width=Inches(6))
                    except Exception:
                        pass

    async def find_title(tags_array):
        for item in tags_array:
            if item[0] == 'TITLE':
                return item[1]

    reply_array = await split_tags(answer)
    await parse_response(reply_array)
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_title = f"{await find_title(reply_array)}.docx"
    print(f"done {docx_title}")

    return docx_bytes, docx_title
